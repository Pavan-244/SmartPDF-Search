import os
import uuid
import warnings
from contextlib import asynccontextmanager
import markdown
from markdown.extensions.codehilite import CodeHiliteExtension
from markdown.extensions.fenced_code import FencedCodeExtension
from markdown.extensions.tables import TableExtension
from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Depends, Form
from fastapi.responses import HTMLResponse, JSONResponse, Response, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from langchain.chains import RetrievalQA
from langchain.chains.llm import LLMChain
from langchain.chains.combine_documents.stuff import StuffDocumentsChain
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import PDFPlumberLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, Session
from langchain_community.chat_models import ChatOpenAI
from models import Base, ChatHistory
from pathlib import Path
from typing import Optional

# Suppress LangChain deprecation warnings
warnings.filterwarnings("ignore", category=DeprecationWarning, module="langchain")

# Lifespan event handler (modern approach)
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    print("🦙 LlamaDoc AI - Starting up...")
    print("✅ Database initialized")
    print("✅ Static files mounted")
    print("✅ Ready to process PDFs!")
    yield
    # Shutdown
    print("🔄 LlamaDoc AI - Shutting down gracefully...")

app = FastAPI(title="PDF QA with LangChain & FastAPI", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Audio storage directory
AUDIO_DIR = Path("uploads/audio")
AUDIO_DIR.mkdir(parents=True, exist_ok=True)

# Mount uploads directory for serving audio files
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

# Database setup
DATABASE_URL = os.environ.get("DATABASE_URL", "sqlite:///./history.db")
engine = create_engine(
    DATABASE_URL, 
    connect_args={"check_same_thread": False} if "sqlite" in DATABASE_URL else {}
)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

# Create tables
Base.metadata.create_all(bind=engine)

# Dependency to get DB session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

OPENAI_API_BASE = os.environ.get("OPENAI_API_BASE", "http://localhost:1234/v1")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "lm-studio")

INDEX_STORE = {}

prompt = """
You are a domain expert assistant.
Use the provided context to answer the question clearly and accurately.
If the answer cannot be found in the context, say "The information is not available in the provided context."
Provide a well-structured answer in 3–4 sentences and keep it factual.

Context:
{context}

Question:
{question}

Answer:
"""
QA_CHAIN_PROMPT = PromptTemplate.from_template(prompt)

def convert_markdown_to_html(text: str) -> str:
    """
    Convert markdown text to HTML with support for:
    - Code blocks with syntax highlighting
    - Tables
    - Lists, headers, bold, italic, etc.
    """
    md = markdown.Markdown(
        extensions=[
            'fenced_code',
            'codehilite',
            'tables',
            'nl2br',  # Convert newlines to <br>
            'sane_lists',
        ],
        extension_configs={
            'codehilite': {
                'css_class': 'highlight',
                'linenums': False,
                'guess_lang': True,
            }
        }
    )
    return md.convert(text)

def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'

@app.get("/favicon.ico")
async def favicon():
    """Serve a simple SVG favicon with a llama emoji"""
    svg_content = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100">
        <text y="80" font-size="80">🦙</text>
    </svg>"""
    return Response(content=svg_content, media_type="image/svg+xml")

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/upload")
async def upload_pdf(file: UploadFile = File(...)):
    if not allowed_file(file.filename):
        raise HTTPException(status_code=400, detail="Invalid file type")

    uid = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_FOLDER, f"{uid}_{file.filename}")
    with open(file_path, "wb") as f:
        f.write(await file.read())

    loader = PDFPlumberLoader(file_path)
    docs = loader.load()

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    documents = text_splitter.split_documents(docs)

    # If no text was extracted, return a clear error to the client
    if not documents:
        raise HTTPException(status_code=400, detail="No text could be extracted from the uploaded PDF")

    # Explicitly provide a model_name to avoid LangChain deprecation warnings
    # and to ensure embeddings can be produced. You can change this model to one
    # available in your environment if needed.
    embedder = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    try:
        vector = FAISS.from_documents(documents, embedder)
    except Exception as e:
        # Catch embedding/index creation errors and return a helpful message
        raise HTTPException(status_code=500, detail=f"Failed to create vector store: {e}")
    retriever = vector.as_retriever(search_type="similarity", search_kwargs={"k": 3})

    INDEX_STORE[uid] = {"retriever": retriever, "vector": vector}

    return JSONResponse({"upload_id": uid, "message": "Index created"})

@app.post("/query")
async def query(request: Request):
    """Accept either JSON or form-encoded POSTs for the query.

    JSON example: {"upload_id": "...", "question": "..."}
    Form example: upload_id and question fields in form data
    """
    upload_id = None
    question = None
    from_form = False

    # Accept JSON body
    try:
        data = await request.json()
        if isinstance(data, dict):
            upload_id = data.get("upload_id")
            question = data.get("question")
    except Exception:
        # Not JSON, try form data
        try:
            form = await request.form()
            upload_id = form.get("upload_id")
            question = form.get("question")
            from_form = True
        except Exception:
            pass

    if not upload_id or not question:
        raise HTTPException(status_code=400, detail="upload_id and question required")
    if upload_id not in INDEX_STORE:
        raise HTTPException(status_code=404, detail="unknown upload_id")

    retriever = INDEX_STORE[upload_id]['retriever']

    llm = ChatOpenAI(
        model="tinyllama-1.1b-chat-v1.0",
        temperature=0.0,
        openai_api_base=OPENAI_API_BASE,
        openai_api_key=OPENAI_API_KEY,
        request_timeout=60,
    )

    llm_chain = LLMChain(llm=llm, prompt=QA_CHAIN_PROMPT)

    document_prompt = PromptTemplate(
        input_variables=["page_content", "source"],
        template="Context:\ncontent:{page_content}\nsource:{source}",
    )

    combine_documents_chain = StuffDocumentsChain(
        llm_chain=llm_chain,
        document_variable_name="context",
        document_prompt=document_prompt,
    )

    qa = RetrievalQA(
        combine_documents_chain=combine_documents_chain,
        retriever=retriever,
        return_source_documents=True,
    )

    # Call the RetrievalQA with a query dict so it returns structured data
    result = qa({"query": question})
    answer = result.get('result')
    sources = []
    for sd in result.get('source_documents', []):
        sources.append({'page_content': sd.page_content[:800], 'metadata': getattr(sd, 'metadata', {})})

    # Build a structured answer payload to make UI formatting easier
    # - text: original answer string
    # - html: markdown converted to HTML
    # - summary: first non-empty paragraph (useful as short summary)
    # - paragraphs: answer split into paragraphs
    if isinstance(answer, str):
        paragraphs = [p.strip() for p in answer.split('\n\n') if p.strip()]
        summary = paragraphs[0] if paragraphs else answer.strip()
        # Convert markdown to HTML for rich rendering
        html_content = convert_markdown_to_html(answer)
    else:
        # Fallback if answer is not a string
        paragraphs = [str(answer)]
        summary = str(answer)
        html_content = convert_markdown_to_html(str(answer))

    formatted_answer = {
        'text': answer,
        'html': html_content,  # Rich HTML content
        'summary': summary,
        'paragraphs': paragraphs,
    }

    # If the request came from a form POST, render the HTML template with the answer
    if from_form:
        return templates.TemplateResponse(
            "index.html",
            {
                "request": request,
                "upload_id": upload_id,
                "answer": formatted_answer,
                "sources": sources,
                "question": question,
            },
        )
    
    # Save to history automatically (for regular text queries)
    try:
        db: Session = next(get_db())
        history_record = ChatHistory(
            upload_id=upload_id,
            question=question,
            answer=answer if isinstance(answer, str) else str(answer),
            summary=summary,
            voice_type="default",
            audio_speed=180,
            is_muted=False
        )
        db.add(history_record)
        db.commit()
    except Exception as e:
        print(f"Warning: Could not save history: {e}")

    return JSONResponse({'answer': formatted_answer, 'sources': sources})

@app.post("/save_history")
async def save_history(
    upload_id: str = Form(None),
    question: str = Form(None),
    answer: str = Form(None),
    summary: str = Form(None),
    question_audio: UploadFile = File(None),
    voice_type: str = Form("default"),
    audio_speed: int = Form(180),
    is_muted: bool = Form(False),
    db: Session = Depends(get_db)
):
    """
    Save chat history entry with optional audio files and voice settings
    Accepts form-data or JSON
    """
    # Save question audio file if present
    q_audio_path = None
    if question_audio and question_audio.filename:
        ext = question_audio.filename.split('.')[-1]
        fname = f"{uuid.uuid4()}.{ext}"
        fp = AUDIO_DIR / fname
        with open(fp, "wb") as f:
            content = await question_audio.read()
            f.write(content)
        q_audio_path = f"/uploads/audio/{fname}"
    
    # Create short summary if not provided
    if not summary and answer:
        # Take first paragraph or first 400 chars
        summary = (answer or "").split("\n\n")[0][:400]
    
    # Create DB record with voice modulation settings
    record = ChatHistory(
        upload_id=upload_id,
        question=question or "",
        answer=answer or "",
        summary=summary or "",
        question_audio_path=q_audio_path,
        answer_audio_path=None,  # Can be populated by /tts endpoint
        voice_type=voice_type or "default",
        audio_speed=audio_speed or 180,
        is_muted=is_muted or False
    )
    
    db.add(record)
    db.commit()
    db.refresh(record)
    
    return JSONResponse({"history_id": record.id, "message": "History saved"})


@app.get("/history")
async def get_history(
    upload_id: Optional[str] = None,
    limit: int = 50,
    db: Session = Depends(get_db)
):
    """
    Retrieve chat history for a given upload_id
    Returns list of history entries ordered by most recent first
    """
    try:
        query = db.query(ChatHistory)
        
        if upload_id:
            query = query.filter(ChatHistory.upload_id == upload_id)
        
        history_entries = query.order_by(ChatHistory.created_at.desc()).limit(limit).all()
        
        return JSONResponse([entry.to_dict() for entry in history_entries])
    except Exception as e:
        # Log the error and return empty list instead of failing
        print(f"Error loading history: {e}")
        return JSONResponse([])

@app.post("/tts")
async def text_to_speech(request: Request, db: Session = Depends(get_db)):
    """
    Server-side TTS using pyttsx3 (offline)
    Supports voice modulation (type, rate, volume)
    """
    try:
        import pyttsx3
        
        body = await request.json()
        text = body.get("text", "")
        voice_type = body.get("voice_type", "default")
        rate = body.get("rate", 180)
        volume = body.get("volume", 1.0)
        
        if not text:
            raise HTTPException(status_code=400, detail="Text is required")
        
        engine = pyttsx3.init()
        voices = engine.getProperty("voices")
        
        # Set voice type
        if voice_type == "female" and len(voices) > 1:
            engine.setProperty("voice", voices[1].id)
        elif voice_type == "male" and len(voices) > 0:
            engine.setProperty("voice", voices[0].id)
        else:
            engine.setProperty("voice", voices[0].id)
        
        # Set speech rate and volume
        engine.setProperty("rate", rate)
        engine.setProperty("volume", volume)
        
        # Generate audio file
        fname = f"{uuid.uuid4()}.mp3"
        path = AUDIO_DIR / fname
        
        engine.save_to_file(text, str(path))
        engine.runAndWait()
        
        audio_url = f"/uploads/audio/{fname}"
        
        return JSONResponse({
            "audio_url": audio_url,
            "message": "TTS audio generated with pyttsx3",
            "settings": {
                "voice_type": voice_type,
                "rate": rate,
                "volume": volume
            }
        })
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="pyttsx3 not installed. Install with: pip install pyttsx3"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"TTS error: {str(e)}")


@app.post("/voice-input")
async def voice_input(audio: UploadFile = File(...)):
    """
    Speech-to-Text endpoint using SpeechRecognition
    Accepts audio file and returns transcribed text
    Supports multiple audio formats (webm, wav, ogg, mp3)
    """
    temp_input = None
    temp_wav = None
    
    try:
        import speech_recognition as sr
        from pydub import AudioSegment
        import time
        
        # Save uploaded audio temporarily
        file_ext = audio.filename.split('.')[-1] if audio.filename else 'webm'
        temp_input = AUDIO_DIR / f"temp_input_{uuid.uuid4()}.{file_ext}"
        temp_wav = AUDIO_DIR / f"temp_converted_{uuid.uuid4()}.wav"
        
        # Save the uploaded file
        content = await audio.read()
        with open(temp_input, "wb") as f:
            f.write(content)
        
        # Ensure file is closed before proceeding
        time.sleep(0.1)
        
        # Convert to WAV format if needed
        audio_path = temp_input
        if file_ext.lower() != 'wav':
            try:
                audio_segment = AudioSegment.from_file(str(temp_input))
                audio_segment.export(str(temp_wav), format='wav')
                audio_path = temp_wav
                # Ensure export is complete
                time.sleep(0.1)
            except Exception as conv_error:
                # If conversion fails, try using browser fallback
                raise HTTPException(
                    status_code=501,
                    detail=f"Audio conversion failed. FFmpeg may not be installed. Error: {str(conv_error)}"
                )
        
        # Initialize recognizer
        recognizer = sr.Recognizer()
        
        # Adjust for ambient noise and transcribe
        transcribed_text = None
        with sr.AudioFile(str(audio_path)) as source:
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            audio_data = recognizer.record(source)
            transcribed_text = recognizer.recognize_google(audio_data)
        
        # Return result before cleanup
        result = JSONResponse({
            "text": transcribed_text,
            "message": "Audio transcribed successfully"
        })
        
        # Clean up temp files with retry logic
        time.sleep(0.2)  # Wait for files to be fully released
        
        for filepath in [temp_input, temp_wav]:
            if filepath and filepath.exists():
                try:
                    os.remove(filepath)
                except PermissionError:
                    # File still locked, try again after a delay
                    time.sleep(0.3)
                    try:
                        os.remove(filepath)
                    except Exception:
                        # If still can't delete, log and continue
                        print(f"Warning: Could not delete temp file: {filepath}")
        
        return result
                
    except sr.UnknownValueError:
        # Clean up before raising
        cleanup_temp_files(temp_input, temp_wav)
        raise HTTPException(
            status_code=400,
            detail="Could not understand audio. Please speak clearly and try again."
        )
    except sr.RequestError as e:
        cleanup_temp_files(temp_input, temp_wav)
        raise HTTPException(
            status_code=503,
            detail=f"Speech recognition service error: {str(e)}"
        )
    except ImportError as e:
        cleanup_temp_files(temp_input, temp_wav)
        if 'speech_recognition' in str(e).lower():
            raise HTTPException(
                status_code=501,
                detail="SpeechRecognition not installed. Install with: pip install SpeechRecognition"
            )
        elif 'pydub' in str(e).lower():
            raise HTTPException(
                status_code=501,
                detail="pydub not installed. Install with: pip install pydub"
            )
        else:
            raise HTTPException(status_code=501, detail=f"Missing dependency: {str(e)}")
    except HTTPException:
        # Re-raise HTTPException without wrapping
        cleanup_temp_files(temp_input, temp_wav)
        raise
    except Exception as e:
        cleanup_temp_files(temp_input, temp_wav)
        raise HTTPException(status_code=500, detail=f"Voice input error: {str(e)}")


def cleanup_temp_files(temp_input, temp_wav):
    """Helper function to clean up temporary audio files"""
    import time
    
    time.sleep(0.2)  # Give system time to release file handles
    
    for filepath in [temp_input, temp_wav]:
        if filepath and filepath.exists():
            try:
                os.remove(filepath)
            except Exception as e:
                print(f"Warning: Could not delete temp file {filepath}: {e}")
                # Try once more after a longer delay
                try:
                    time.sleep(0.5)
                    os.remove(filepath)
                except Exception:
                    pass  # Give up gracefully


@app.get("/download/{upload_id}/{format}")
async def download_answer(upload_id: str, format: str, db: Session = Depends(get_db)):
    """
    Download the latest answer for a PDF in TXT, PDF, or DOCX format
    """
    # Get the latest answer for this upload_id
    history = db.query(ChatHistory).filter(
        ChatHistory.upload_id == upload_id
    ).order_by(ChatHistory.created_at.desc()).first()
    
    if not history:
        raise HTTPException(status_code=404, detail="No answers found for this PDF")
    
    text = f"Question: {history.question}\n\nAnswer:\n{history.answer}"
    
    # Create downloads directory
    downloads_dir = Path("downloads")
    downloads_dir.mkdir(exist_ok=True)
    
    filename = f"answer_{upload_id[:8]}"
    
    try:
        if format == "txt":
            filepath = downloads_dir / f"{filename}.txt"
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(text)
        
        elif format == "docx":
            from docx import Document
            filepath = downloads_dir / f"{filename}.docx"
            doc = Document()
            doc.add_heading("AI Generated Answer", level=1)
            doc.add_heading("Question:", level=2)
            doc.add_paragraph(history.question)
            doc.add_heading("Answer:", level=2)
            doc.add_paragraph(history.answer)
            doc.save(str(filepath))
        
        elif format == "pdf":
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import simpleSplit
            
            filepath = downloads_dir / f"{filename}.pdf"
            c = canvas.Canvas(str(filepath), pagesize=letter)
            width, height = letter
            
            # Title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, height - 50, "AI Generated Answer")
            
            # Question
            y = height - 100
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, "Question:")
            y -= 20
            
            c.setFont("Helvetica", 10)
            lines = simpleSplit(history.question, "Helvetica", 10, width - 100)
            for line in lines:
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line)
                y -= 15
            
            # Answer
            y -= 20
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, "Answer:")
            y -= 20
            
            c.setFont("Helvetica", 10)
            lines = simpleSplit(history.answer, "Helvetica", 10, width - 100)
            for line in lines:
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line)
                y -= 15
            
            c.save()
        
        else:
            raise HTTPException(status_code=400, detail="Unsupported format. Use txt, pdf, or docx")
        
        return FileResponse(
            path=str(filepath),
            filename=f"{filename}.{format}",
            media_type="application/octet-stream"
        )
    
    except ImportError as e:
        raise HTTPException(
            status_code=501,
            detail=f"Required library not installed: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download error: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)